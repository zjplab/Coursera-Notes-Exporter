from bs4 import BeautifulSoup
import docx
import urllib.request
import os
from selenium import webdriver
from selenium.webdriver.firefox.service import Service
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
driver_path="geckodriver.exe"
firefox_services = Service(executable_path=driver_path, port=3000, service_args=['--marionette-port', '2828', '--connect-existing'])
driver = webdriver.Firefox(service=firefox_services)
driver.get('https://www.coursera.org/learn/devops-aws-code-build-test/lecture/qAv4B/thinking-in-devops')
# html = driver.page_source
# wait until entire document is ready
WebDriverWait(driver, 100).until(EC.presence_of_element_located((By.CSS_SELECTOR, ".ItemLecture_Video_Notes_Navigation > div:nth-child(1) > button:nth-child(1)")))
notes_button=driver.find_element(By.CSS_SELECTOR, ".ItemLecture_Video_Notes_Navigation > div:nth-child(1) > button:nth-child(1)")
notes_button.click()

# WebDriverWait(driver,10).until(EC.element_to_be_clickable((By.CSS_SELECTOR,"button[aria-controls='video-notes-drawer']"))).click()
time.sleep(5)
innerHTML = driver.execute_script("return document.body.innerHTML")
soup = BeautifulSoup(innerHTML, "lxml")
title=soup.select('h1[class*="video-name"]')[0].get_text()
print(title)
doc = docx.Document()

# Add a Title to the document
doc.add_heading(title, 0)
# function for downloading an image from a url with max retry of 3 times
def download_file(url, file_name, max_retry=3):
    try:
        urllib.request.urlretrieve(url, file_name)
    except Exception as e:
        print(e)
        if max_retry > 0:
            download_file(url, file_name, max_retry - 1)
        else:
            print('Failed to download file after 3 retries')

for img_tag in soup.find_all('img', {'alt': 'Video thumbnail'}):
    # download image from img url, save to `video.png`
    print(img_tag['src'])
    download_file(img_tag['src'], file_name='video.png')
    doc.add_picture('video.png')

# # Now save the document to a location
doc.save('export.docx')
os.remove('video.png')